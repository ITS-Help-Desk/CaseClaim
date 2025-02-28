from discord import app_commands
from discord.ext import commands
import discord
import traceback
from bot.models.checked_claim import CheckedClaim
from bot.models.user import User
from bot.status import Status
import statistics
import datetime
from docx import Document
from docx.shared import Pt
import zipfile

# Use TYPE_CHECKING to avoid circular import from bot
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from ..bot import Bot


class GenEvalCommand(commands.Cog):
    def __init__(self, bot: "Bot") -> None:
        """Creates the /geneval command using a cog.

        Args:
            bot (Bot): A reference to the original Bot instantiation.
        """
        self.bot = bot

    @app_commands.command(description="Automatically generates monthly evals")
    async def geneval(self, interaction: discord.Interaction,  month: int, year: int) -> None:
        """

        Args:
            interaction (discord.Interaction): Interaction that the slash command originated from
        """
        total_hd_cases, total_checked_cases, total_pinged_cases, total_kudos_cases = self.get_data(month,year)

        hd_total_claims, median_claim, median_ping_percent, top_claim_percent, data = self.organize_data_for_word(total_hd_cases, total_checked_cases, total_pinged_cases, total_kudos_cases)

        now = datetime.datetime.now()
        filenames = []

        for user_id in list(data.keys()):
            user = User.from_id(self.bot.connection, user_id)

            discord_user = await interaction.guild.fetch_member(user_id)

            if self.bot.check_if_lead(discord_user):
                continue

            fields, template_name = self.create_word_fields(
                user.full_name, 
                now,
                hd_total_claims,
                median_claim,
                median_ping_percent,
                top_claim_percent,
                data[user_id]
            )

            title = f"evals/{user.full_name.split(' ')[-1]} {now.strftime('%B').upper()} {now.strftime('%Y')}.docx"
            filenames.append(title)
            self.create_document(template_name, title, fields)
        
        zipname = f"evals/eval{month}{year}.zip"
        self.create_zip(filenames, zipname)

        await interaction.response.send_message(content="hello", file=discord.File(zipname), ephemeral=True, delete_after=300)

    def create_zip(self, filenames: list[str], zipname: str):
        with zipfile.ZipFile(zipname, 'w') as zipf:
            for file in filenames:
                zipf.write(file)

    def create_document(self, template: str, save_title: str, fields: dict[int, str]):
        document = Document(template)
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        i = 1
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        #print(i, paragraph.text)
                        if i in list(fields.keys()):
                            paragraph.text = fields[i]
                            paragraph.style = document.styles['Normal']
                        i += 1

        document.save(save_title)
    
    def create_word_fields(self, username, date, hd_total_claims, median_claim, median_ping_percent, top_claim_percent, data):
        kudos = data["kudos"]
        pings = data["pings"]

        fields = {
            18: username,
            20: 0,

            31: date.strftime("%m/%d/%Y"),

            37: kudos[0] if len(kudos) > 0 else "",
            38: kudos[1] if len(kudos) > 1 else "",
            39: kudos[2] if len(kudos) > 2 else "",
            40: kudos[3] if len(kudos) > 3 else "",
            41: kudos[4] if len(kudos) > 4 else "",

            52: pings[0] if len(pings) > 0 else "",
            53: pings[1] if len(pings) > 1 else "",
            54: pings[2] if len(pings) > 2 else "",
            55: pings[3] if len(pings) > 3 else "",
            56: pings[4] if len(pings) > 4 else "",
        }
        
        if len(pings) < 6:
            template_name = "templates/template1row.docx"
            offset = 0
        
        if len(pings) > 5:
            template_name = "templates/template2row.docx"
            offset = 5
            fields[57] = pings[5] if len(pings) > 5 else ""
            fields[58] = pings[6] if len(pings) > 6 else ""
            fields[59] = pings[7] if len(pings) > 7 else ""
            fields[60] = pings[8] if len(pings) > 8 else ""
            fields[61] = pings[9] if len(pings) > 9 else ""
        
        if len(pings) > 10:
            template_name = "templates/template3row.docx"
            offset = 10
            fields[62] = pings[10] if len(pings) > 10 else ""
            fields[63] = pings[11] if len(pings) > 11 else ""
            fields[64] = pings[12] if len(pings) > 12 else ""
            fields[65] = pings[13] if len(pings) > 13 else ""
            fields[66] = pings[14] if len(pings) > 14 else ""
        if len(pings) > 15:
            template_name = "templates/template4row.docx"
            offset = 15
            fields[67] = pings[15] if len(pings) > 15 else ""
            fields[68] = pings[16] if len(pings) > 16 else ""
            fields[69] = pings[17] if len(pings) > 17 else ""
            fields[70] = pings[18] if len(pings) > 18 else ""
            fields[71] = pings[19] if len(pings) > 19 else ""
        
        if len(pings) > 20:
            template_name = "templates/template5row.docx"
            offset = 20
            fields[72] = pings[20] if len(pings) > 20 else ""
            fields[73] = pings[21] if len(pings) > 21 else ""
            fields[74] = pings[22] if len(pings) > 22 else ""
            fields[75] = pings[23] if len(pings) > 23 else ""
            fields[76] = pings[24] if len(pings) > 24 else ""
        
        # Top stats label
        fields[62 + offset] = f"{data['user_pinged_count']} / {data['user_total_count']} / {hd_total_claims}"

        # User case stats
        fields[65 + offset] = f"Individual: {data['claim_percent']:.2%}"
        fields[66 + offset] = f"Team Median: {median_claim:.2%}"
        fields[67 + offset] = f"Team Top: {top_claim_percent:.2%}"

        # Ping stats
        fields[70 + offset] = f"Individual: {data['ping_percent']:.2%}"
        fields[71 + offset] = f"Team Median: {median_ping_percent:.2%}"

        return fields, template_name

    def organize_data_for_word(self, total_case_count: int, checked: dict[int, int], pinged: dict[int, list[str]], kudos: dict[int, list[str]]):
        organized_data = {}

        claim_median_list = []
        top_claim = 0
        ping_median_list = []

        for user_id in list(checked.keys()):
            claim_count = len(pinged[user_id]) + len(kudos[user_id]) + checked[user_id]
            if claim_count > top_claim:
                top_claim = claim_count
            claim_median_list.append(claim_count / total_case_count)
            ping_median_list.append(len(pinged[user_id]) / claim_count)

            organized_data[user_id] = {
                "user_pinged_count": len(pinged[user_id]),
                "user_total_count": claim_count,

                "claim_percent": claim_count / total_case_count,

                "pings": pinged[user_id],
                "kudos": kudos[user_id],
                
                "ping_percent": len(pinged[user_id]) / claim_count
            }

        median_claim = int(statistics.median(claim_median_list))
        median_ping = statistics.mean(ping_median_list)

        return total_case_count, median_claim, median_ping, top_claim / total_case_count, organized_data

    def get_data(self, month: int, year: int) -> tuple[int, dict[int, int], dict[int, list[str]], dict[int, list[str]]]:
        all_cases = CheckedClaim.get_all_from_month(self.bot.connection, month, year)

        # Tech data
        total_hd_cases = 0
        total_checked_cases = {}
        total_pinged_cases = {}
        total_kudos_cases = {}

        for case in all_cases:
            if case.status == str(Status.DONE):
                continue

            total_hd_cases += 1
            
            # Initialize tech data
            total_checked_cases.setdefault(case.tech.discord_id, 0)
            total_pinged_cases.setdefault(case.tech.discord_id, [])
            total_kudos_cases.setdefault(case.tech.discord_id, [])

            # Update tech data
            if case.status == str(Status.CHECKED):
                total_checked_cases[case.tech.discord_id] += 1
            elif case.status == str(Status.PINGED) or case.status == str(Status.RESOLVED):
                total_pinged_cases[case.tech.discord_id].append(case.case_num)
            elif case.status == str(Status.KUDOS):
                total_kudos_cases[case.tech.discord_id].append(case.case_num)

        return total_hd_cases, total_checked_cases, total_pinged_cases, total_kudos_cases

    @geneval.error
    async def geneval_error(self, ctx: discord.Interaction, error):
        full_error = traceback.format_exc()

        ch = await self.bot.fetch_channel(self.bot.error_channel)

        msg = f"Error with **/geneval** ran by <@!{ctx.user.id}>.\n```{full_error}```"
        if len(msg) > 1993:
            msg = msg[:1993] + "...```"
        await ch.send(msg)
